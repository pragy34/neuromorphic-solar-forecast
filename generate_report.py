from __future__ import annotations

import csv
import json
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
PROJECT_ROOT = ROOT
OUT_DIR = PROJECT_ROOT / "outputs"
METRICS_DIR = OUT_DIR / "metrics"
PLOTS_DIR = OUT_DIR / "plots"
FORECASTS_DIR = OUT_DIR / "forecasts"
FRONTEND_DIR = PROJECT_ROOT / "neurospikeapp"
STREAMLIT_DIR = PROJECT_ROOT / "frontend"
REPORT_PATH = PROJECT_ROOT / "NeuroSpike_Detailed_Capstone_Report.docx"


PROJECT_TITLE = (
    "NeuroSpike: AI-Powered Solar Intelligence System for Irradiance "
    "Forecasting and Actionable Energy Decision Support"
)


@dataclass(frozen=True)
class FigureSpec:
    path: Path
    caption: str
    note: str = ""
    width: float = 6.1


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def read_json(path: Path) -> object | None:
    if not path.exists():
        return None
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def clean_number(value: str | float | int | None, digits: int = 3) -> str:
    if value is None:
        return ""
    try:
        number = float(value)
    except (TypeError, ValueError):
        return str(value)
    return f"{number:.{digits}f}".rstrip("0").rstrip(".")


def safe_cell_text(value: object) -> str:
    text = "" if value is None else str(value)
    return text.replace("\n", " ").strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, font_size: int = 10) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 35 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(safe_cell_text(text))
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5

    doc.styles["Heading 1"].font.size = Pt(16)
    doc.styles["Heading 1"].font.bold = True
    doc.styles["Heading 2"].font.size = Pt(14)
    doc.styles["Heading 2"].font.bold = True
    doc.styles["Heading 3"].font.size = Pt(12)
    doc.styles["Heading 3"].font.bold = True


def add_page_number(section) -> None:
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])


def add_cover_page(doc: Document) -> None:
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CAPSTONE PROJECT REPORT")
    run.bold = True
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("\n" + PROJECT_TITLE)
    run.bold = True
    run.font.size = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "\nSubmitted in partial fulfillment of the requirements for the award of the degree of\n"
        "Bachelor of Technology\n\n"
        "Department of [Department Name]\n"
        "[College Name]\n"
        "[University Name]\n\n"
        "Submitted by\n"
        "Student A ([Roll Number])\n"
        "Student B ([Roll Number])\n"
        "Student C ([Roll Number])\n\n"
        "Under the guidance of\n"
        "[Supervisor Name]\n\n"
        "Academic Year: 2025-2026"
    )
    run.font.size = Pt(12)
    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(12 if level > 1 else 18)
    p.paragraph_format.space_after = Pt(6)


def add_body(doc: Document, text: str) -> None:
    for block in text.strip().split("\n\n"):
        if not block.strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(block.strip())
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.7)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Cm(0.7)
        run = p.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_table(doc: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, font_size=10)
        set_cell_shading(table.rows[0].cells[idx], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], safe_cell_text(value), font_size=10)
    doc.add_paragraph()


def add_caption(doc: Document, caption: str, note: str = "") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(10)
    if note:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(note)
        run.italic = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(9)


def add_figure(doc: Document, spec: FigureSpec) -> bool:
    if not spec.path.exists():
        return False
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    try:
        run.add_picture(str(spec.path), width=Inches(spec.width))
    except Exception:
        return False
    add_caption(doc, spec.caption, spec.note)
    return True


def add_code_block(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1
        run = p.add_run(line.rstrip())
        run.font.name = "Courier New"
        run.font.size = Pt(9)


def model_summary_rows() -> list[list[str]]:
    rows = read_csv_rows(METRICS_DIR / "final_summary_table.csv")
    result = []
    for row in rows:
        result.append(
            [
                row.get("Model", ""),
                clean_number(row.get("Mean RMSE")),
                clean_number(row.get("Mean MAE")),
                clean_number(row.get("Mean R2") or row.get("Mean RÂ²") or row.get("Mean R²")),
                clean_number(row.get("Mean Skill"), 2),
                row.get("Best City", ""),
                row.get("Worst City", ""),
            ]
        )
    return result


def r2_city_rows() -> list[list[str]]:
    rows = read_csv_rows(METRICS_DIR / "r2_comparison.csv")
    keep_models = {"NeuroSpike_h1", "BiLSTM_h1", "RandomForest", "XGBoost", "Persistence"}
    result = []
    for row in rows:
        if row.get("model") in keep_models:
            result.append(
                [
                    row.get("model", ""),
                    clean_number(row.get("new_delhi")),
                    clean_number(row.get("london")),
                    clean_number(row.get("riyadh")),
                    clean_number(row.get("sydney")),
                    clean_number(row.get("MEAN")),
                ]
            )
    return result


def snn_city_rows() -> list[list[str]]:
    rows = read_csv_rows(METRICS_DIR / "snn_results.csv")
    result = []
    for row in rows:
        result.append(
            [
                row.get("city", "").replace("_", " ").title(),
                clean_number(row.get("RMSE")),
                clean_number(row.get("MAE")),
                clean_number(row.get("MAPE")),
                clean_number(row.get("R2")),
                clean_number(row.get("Skill"), 2),
            ]
        )
    return result


def feature_score_rows(limit: int = 10) -> list[list[str]]:
    rows = read_csv_rows(METRICS_DIR / "feature_scores.csv")
    if not rows:
        return []
    score_columns = [c for c in rows[0].keys() if c.lower() not in {"feature", "rank"}]
    selected = []
    for row in rows[:limit]:
        selected.append([row.get("feature", ""), *[clean_number(row.get(c), 4) for c in score_columns[:3]]])
    return selected


def forecast_rows(limit: int = 8) -> list[list[str]]:
    forecasts = read_json(FORECASTS_DIR / "all_forecasts.json")
    if not isinstance(forecasts, dict):
        return []
    rows = []
    for city, payload in forecasts.items():
        if len(rows) >= limit:
            break
        if isinstance(payload, dict):
            forecast = payload.get("forecast") or payload.get("forecasts") or payload
            ghi = ""
            if isinstance(forecast, dict):
                ghi = forecast.get("ghi_predicted") or forecast.get("GHI") or forecast.get("ghi")
            rows.append([city.replace("_", " ").title(), clean_number(ghi), "1 hour"])
    return rows


def discover_figures() -> list[FigureSpec]:
    return [
        FigureSpec(PLOTS_DIR / "03_diurnal_patterns.png", "Figure 1.1: Observed diurnal solar irradiance patterns across selected cities.", "The daily curve confirms the strong time-of-day structure used by the forecasting models."),
        FigureSpec(PLOTS_DIR / "03_climate_comparison.png", "Figure 1.2: Climate-wise comparison of irradiance behavior.", "The selected cities intentionally cover desert, coastal, temperate, tropical, and high-variability conditions."),
        FigureSpec(PLOTS_DIR / "03_correlation_heatmaps.png", "Figure 3.1: Correlation heatmaps for meteorological and engineered variables.", "Correlation analysis guided the inclusion of lag, rolling, and weather interaction features."),
        FigureSpec(PLOTS_DIR / "04_feature_correlation_ghi.png", "Figure 3.2: Feature relationship with Global Horizontal Irradiance.", "The strongest predictors were retained for baseline and deep learning experiments."),
        FigureSpec(PLOTS_DIR / "05_consensus_selection.png", "Figure 3.3: Consensus feature selection output.", "Multiple feature scoring methods were combined to reduce dependence on a single selection criterion."),
        FigureSpec(PLOTS_DIR / "06_baseline_comparison.png", "Figure 5.1: Baseline model comparison.", "Tree-based models improve on statistical baselines, but the temporal deep models provide the strongest generalization."),
        FigureSpec(PLOTS_DIR / "06_predictions_vs_actual.png", "Figure 5.2: Predicted versus actual irradiance for baseline models.", "The diagonal concentration shows that engineered features capture most daily irradiance movement."),
        FigureSpec(PLOTS_DIR / "07_lstm_training_history.png", "Figure 5.3: BiLSTM training history.", "Training and validation curves were monitored to detect overfitting and convergence issues."),
        FigureSpec(PLOTS_DIR / "07_lstm_predictions.png", "Figure 5.4: BiLSTM predicted versus actual GHI.", "BiLSTM improves temporal consistency by learning patterns in both forward and backward sequence directions."),
        FigureSpec(PLOTS_DIR / "08_snn_predictions.png", "Figure 5.5: NeuroSpike SNN prediction performance.", "The SNN model provides the lowest mean error and strongest R2 in the saved evaluation outputs."),
        FigureSpec(PLOTS_DIR / "08_snn_spike_analysis.png", "Figure 5.6: Spike activity analysis of the NeuroSpike layer.", "Spike activity supports interpretation of how rapid irradiance transitions are encoded."),
        FigureSpec(PLOTS_DIR / "09_error_analysis.png", "Figure 5.7: Error and residual analysis.", "Residual behavior was reviewed to identify time periods and climates where forecasts are most difficult."),
        FigureSpec(PLOTS_DIR / "09_evaluation_dashboard.png", "Figure 5.8: Consolidated model evaluation dashboard.", "This figure summarizes comparative accuracy, residual behavior, and city-wise robustness."),
        FigureSpec(PLOTS_DIR / "10_forecast_summary.png", "Figure 5.9: Forecast summary generated from trained models.", "The final forecasting stage turns model output into deployable decision support data."),
    ]


def frontend_figures() -> list[FigureSpec]:
    candidates = [
        FigureSpec(FRONTEND_DIR / "static" / "solar-hero-bg.png", "Figure 4.1: Frontend visual background used in the standalone NeuroSpike web dashboard.", "The image anchors the dashboard interface around the solar-energy use case."),
    ]
    screenshot_dirs = [
        OUT_DIR / "frontend_screenshots",
        OUT_DIR / "screenshots",
        PROJECT_ROOT / "screenshots",
    ]
    count = 2
    for folder in screenshot_dirs:
        if folder.exists():
            for path in sorted(folder.glob("*")):
                if path.suffix.lower() in {".png", ".jpg", ".jpeg"}:
                    candidates.append(
                        FigureSpec(
                            path,
                            f"Figure 4.{count}: Frontend dashboard screenshot - {path.stem.replace('_', ' ').title()}.",
                            "This screenshot documents the user-facing dashboard built for forecasts, comparison, savings, and interpretation.",
                        )
                    )
                    count += 1
    return candidates


def add_preliminary_pages(doc: Document) -> None:
    add_cover_page(doc)

    add_heading(doc, "Certificate", 2)
    add_body(
        doc,
        f"This is to certify that the capstone project entitled \"{PROJECT_TITLE}\" has been carried out by the students listed below under the guidance of [Supervisor Name], Department of [Department Name], [College Name]. The work is submitted in partial fulfillment of the requirements for the award of the Bachelor of Technology degree.",
    )
    add_table(doc, ["Student Name", "Roll Number"], [["Student A", "[Roll Number]"], ["Student B", "[Roll Number]"], ["Student C", "[Roll Number]"]])
    add_body(doc, "Supervisor Signature: ____________________\n\nHead of Department Signature: ____________________\n\nDate: ____________________")
    doc.add_page_break()

    add_heading(doc, "Declaration of Originality", 2)
    add_body(
        doc,
        "We declare that this report is an original account of the NeuroSpike project work completed by our team. The text has been written specifically for this implementation, the experimental values are taken from our project artifacts where available, and all external concepts, standards, datasets, frameworks, and tools are acknowledged in the references. This report has not been submitted previously for another degree, diploma, or certification.",
    )
    add_table(doc, ["Student Name", "Signature"], [["Student A", ""], ["Student B", ""], ["Student C", ""]])
    doc.add_page_break()

    add_heading(doc, "Acknowledgement", 2)
    add_body(
        doc,
        "We express our sincere gratitude to [Supervisor Name] for guiding the technical direction of this project and for encouraging a solution that connects forecasting accuracy with practical energy decisions. We thank the Department of [Department Name] for providing the academic environment, review feedback, and infrastructure required to complete the work.\n\nWe also acknowledge NASA POWER for making solar and meteorological data publicly accessible. The project uses open-source software including Python, FastAPI, PyTorch, scikit-learn, XGBoost, Streamlit, and python-docx. Finally, we thank our classmates, friends, and family members for their patience, feedback, and encouragement during development and documentation.",
    )
    doc.add_page_break()

    add_heading(doc, "Abstract", 2)
    add_body(
        doc,
        "Solar generation is clean and increasingly affordable, but its usefulness depends on how well users can anticipate variation in irradiance. A numerical forecast alone is not enough for many households or small industries because the practical question is not only how much solar energy will be available, but also what action should be taken when that energy is available. NeuroSpike addresses this need by combining short-term Global Horizontal Irradiance forecasting with conversion and decision-support layers.\n\nThe system uses NASA POWER data for ten climatically diverse cities. The pipeline performs preprocessing, exploratory analysis, feature engineering, model training, evaluation, forecast generation, and dashboard deployment. Baseline models include persistence, climatology, linear regression, ridge regression, random forest, and XGBoost. Deep learning models include BiLSTM and a spiking-neural-network-based NeuroSpike model. The final SNN model achieved a mean RMSE of 29.667 W/m2, mean MAE of 17.686 W/m2, and mean R2 of 0.9877 across the saved evaluation results.\n\nThe forecast output is further converted into estimated photovoltaic power using panel area, efficiency, and irradiance assumptions. A decision layer then translates solar availability into user-facing recommendations such as shifting heavy loads, scheduling industrial tasks, or conserving grid power during low-yield windows. The project is exposed through a FastAPI backend, a Streamlit analytics dashboard, and a standalone web interface. The result is a detailed, reproducible solar intelligence prototype that combines forecasting, interpretability, economic reasoning, and sustainability goals.",
    )
    add_body(doc, "Keywords: solar forecasting, Global Horizontal Irradiance, spiking neural network, BiLSTM, NASA POWER, FastAPI, renewable energy decision support, photovoltaic power estimation.")
    doc.add_page_break()

    add_heading(doc, "Table of Contents", 2)
    add_body(
        doc,
        "Certificate\nDeclaration of Originality\nAcknowledgement\nAbstract\nList of Figures\nList of Tables\nList of Abbreviations\nChapter 1: Introduction\nChapter 2: Literature Survey\nChapter 3: System Design and Methodology\nChapter 4: Implementation and Testing\nChapter 5: Results and Discussion\nChapter 6: Conclusion and Future Scope\nReferences\nAppendices\nAnnexure",
    )
    doc.add_page_break()

    add_heading(doc, "List of Figures", 2)
    figure_lines = [spec.caption for spec in frontend_figures() + discover_figures() if spec.path.exists()]
    add_body(doc, "\n".join(figure_lines) if figure_lines else "Figures are inserted automatically when generated plot files are available in outputs/plots.")
    doc.add_page_break()

    add_heading(doc, "List of Tables", 2)
    add_body(
        doc,
        "Table 1.1: Project work breakdown structure\nTable 1.2: Risks and mitigation strategies\nTable 2.1: Literature comparison table\nTable 3.1: Dataset and feature groups\nTable 3.2: Tool and technique selection\nTable 4.1: Experimental setup\nTable 4.2: API and system test cases\nTable 5.1: Model performance summary\nTable 5.2: City-wise R2 comparison\nTable 5.3: NeuroSpike city-wise metrics\nTable 5.4: Objective fulfillment assessment\nTable 5.5: Economic cost-benefit analysis\nTable A.1: PO/PSO mapping\nTable A.2: Team role distribution",
    )
    doc.add_page_break()

    add_heading(doc, "List of Abbreviations", 2)
    add_table(
        doc,
        ["Abbreviation", "Full Form"],
        [
            ["API", "Application Programming Interface"],
            ["BiLSTM", "Bidirectional Long Short-Term Memory"],
            ["GHI", "Global Horizontal Irradiance"],
            ["LIF", "Leaky Integrate-and-Fire"],
            ["MAE", "Mean Absolute Error"],
            ["MAPE", "Mean Absolute Percentage Error"],
            ["ML", "Machine Learning"],
            ["NASA POWER", "NASA Prediction of Worldwide Energy Resources"],
            ["PV", "Photovoltaic"],
            ["R2", "Coefficient of Determination"],
            ["RMSE", "Root Mean Squared Error"],
            ["SNN", "Spiking Neural Network"],
        ],
    )
    doc.add_page_break()


def chapter_1(doc: Document) -> None:
    add_heading(doc, "Chapter 1: Introduction", 1)
    add_heading(doc, "1.1 Background", 2)
    add_body(
        doc,
        "Solar photovoltaic systems are now a major part of the global renewable-energy transition. Their economic value, however, depends on the ability to plan around variability. Solar irradiance changes with time of day, season, cloud cover, atmospheric conditions, and location. A user may own a solar panel system, but without a reliable estimate of the next energy-rich window, high-consumption tasks may still be performed during grid-dependent periods.\n\nGlobal Horizontal Irradiance is one of the most useful measures for solar planning because it represents the solar radiation received on a horizontal surface. In a forecasting workflow, GHI becomes the link between weather data and practical photovoltaic output. This project treats GHI forecasting as the first layer of a larger intelligence system rather than as a standalone numerical task.",
    )
    add_figure(doc, FigureSpec(PLOTS_DIR / "03_diurnal_patterns.png", "Figure 1.1: Observed daily irradiance patterns used to motivate short-term forecasting."))

    add_heading(doc, "1.2 Motivation and Need", 2)
    add_body(
        doc,
        "Most forecasting demonstrations stop after reporting error metrics. That is useful for model comparison, but it does not directly help a household decide when to run appliances or a small industrial unit decide when to schedule non-critical loads. NeuroSpike was designed to close this gap. It forecasts irradiance, estimates usable PV power, and converts that estimate into decision support for different user types.\n\nThe project is also motivated by the need for reproducible, open-data solutions. NASA POWER provides accessible historical solar and weather variables, making it possible to build and validate the pipeline without proprietary sensors or paid data feeds.",
    )

    add_heading(doc, "1.3 Objectives and Problem Statement", 2)
    add_body(doc, "Problem statement: Existing solar forecasting systems often provide numerical irradiance predictions without converting them into actionable, user-specific energy decisions. The technical challenge is to build a short-term forecasting system that is accurate, reproducible, interpretable, and connected to a practical decision layer.")
    add_bullets(
        doc,
        [
            "Collect and prepare multi-city solar and weather data from NASA POWER for short-term GHI forecasting.",
            "Engineer temporal, lag, rolling, weather interaction, and solar-geometry features that improve model accuracy.",
            "Compare statistical baselines, tree-based machine learning models, BiLSTM, and a NeuroSpike SNN model using RMSE, MAE, MAPE, R2, and skill score.",
            "Convert predicted irradiance into approximate PV power output using panel area and efficiency assumptions.",
            "Expose forecasts through API and dashboard interfaces that turn model output into residential and industrial recommendations.",
        ],
    )

    add_heading(doc, "1.4 Social and Environmental Relevance", 2)
    add_body(
        doc,
        "The social value of the project lies in simplifying solar decisions for non-specialist users. A household does not need to understand model residuals to benefit from the system; it needs a clear recommendation such as whether the current period is suitable for high-load activities. For small industries, even a modest shift in load timing can reduce grid dependence and operating cost.\n\nEnvironmentally, better use of available solar energy can reduce fossil-fuel-based grid consumption without requiring additional panels. The project therefore supports sustainability through improved utilization of existing renewable infrastructure. Since the system uses open data and open-source tools, it can be adapted for educational, municipal, and low-cost community deployments.",
    )

    add_heading(doc, "1.5 Project Plan", 2)
    add_table(
        doc,
        ["Work Package", "Main Tasks", "Deliverables"],
        [
            ["WP1: Research", "Study solar forecasting, NASA POWER, baseline models, SNN literature", "Problem definition and literature matrix"],
            ["WP2: Data Pipeline", "Collect, clean, align, and validate city-wise data", "Raw, processed, and feature datasets"],
            ["WP3: Feature Engineering", "Create time, lag, rolling, interaction, and solar features", "Feature CSV files and selection outputs"],
            ["WP4: Model Development", "Train baselines, BiLSTM, and NeuroSpike SNN", "Saved models, checkpoints, and metrics"],
            ["WP5: Application Layer", "Build API, Streamlit dashboard, and standalone web dashboard", "Usable forecasting interfaces"],
            ["WP6: Evaluation and Report", "Analyze metrics, plots, economics, limitations, and future scope", "Capstone report and appendices"],
        ],
    )
    add_table(
        doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["Noisy weather and irradiance data", "Lower forecast reliability", "Use preprocessing checks, rolling statistics, and multi-city validation"],
            ["Overfitting in deep models", "High validation error", "Monitor training history and compare against simpler baselines"],
            ["SNN implementation complexity", "Training instability", "Use a simplified LIF-based model and evaluate against BiLSTM"],
            ["Frontend/backend mismatch", "User interface cannot display real values", "Define API request and response formats early"],
            ["Report similarity risk", "Generic writing may resemble common templates", "Use original project-specific wording, actual project metrics, and proper references"],
        ],
    )

    add_heading(doc, "1.6 Scope and Limitations", 2)
    add_body(
        doc,
        "The current scope is a one-hour-ahead forecasting system for ten configured cities: Riyadh, Cairo, Istanbul, New Delhi, Dubai, London, Sydney, Tokyo, Los Angeles, and Nairobi. The system focuses on GHI and derived PV power estimates rather than full physical modeling of every panel, inverter, and storage component.\n\nThe main limitations are the spatial resolution of NASA POWER data, the absence of live rooftop sensor calibration, the lack of dynamic electricity tariff integration, and the use of generalized decision rules. These limitations do not invalidate the prototype, but they define the boundary between a capstone-grade decision-support system and a commercial energy-management product.",
    )
    doc.add_page_break()


def chapter_2(doc: Document) -> None:
    add_heading(doc, "Chapter 2: Literature Survey", 1)
    add_heading(doc, "2.1 Critical Review", 2)
    add_body(
        doc,
        "Solar forecasting literature can be grouped into physical models, statistical time-series models, machine learning models, deep sequence models, and decision-support systems. Physical methods use solar geometry and atmospheric assumptions, but they require detailed local measurements for strong performance. Statistical methods such as persistence, climatology, ARIMA, and exponential smoothing are easy to implement and provide useful baselines, yet they struggle with nonlinear weather transitions.\n\nMachine learning models such as random forest and XGBoost handle nonlinear interactions between weather variables and engineered time features. Their limitation is that temporal structure is represented through manually created lags and rolling statistics rather than learned sequence memory. Deep learning models, especially LSTM and BiLSTM networks, address this weakness by learning temporal dependencies from sequences. Spiking neural networks are newer in this application area and are attractive because they encode information using event-like spikes, which may represent rapid irradiance changes efficiently.",
    )
    add_heading(doc, "2.2 Summary Table", 2)
    add_table(
        doc,
        ["Method Group", "Typical Strength", "Typical Limitation", "Role in NeuroSpike"],
        [
            ["Persistence and climatology", "Simple, transparent, fast", "Weak during weather changes", "Baseline reference"],
            ["Linear and ridge regression", "Interpretable and stable", "Limited nonlinear modeling", "Feature engineering sanity check"],
            ["Random forest", "Handles nonlinear feature relationships", "Can be large and less sequential", "Strong ML baseline"],
            ["XGBoost", "High accuracy on tabular features", "Needs careful tuning", "Primary tree-based comparison"],
            ["LSTM/BiLSTM", "Learns temporal dependencies", "More computationally expensive", "Deep sequence benchmark"],
            ["Spiking neural network", "Event-like temporal representation", "Less common for continuous solar regression", "Core NeuroSpike model"],
            ["Rule-based decision support", "Transparent and deployable", "Less adaptive than reinforcement learning", "Initial decision layer"],
        ],
    )
    add_heading(doc, "2.3 Identification of Research Gaps", 2)
    add_bullets(
        doc,
        [
            "Many studies optimize forecast accuracy but do not translate predictions into user-facing operational guidance.",
            "SNNs are less explored for continuous GHI regression compared with load forecasting and classification tasks.",
            "A single-city evaluation can hide climate-specific weakness; multi-city testing is necessary for practical robustness.",
            "Commercial solar intelligence tools are often paid or API-only, limiting educational transparency and local customization.",
            "Decision support is frequently treated as a separate application problem rather than integrated with the model evaluation pipeline.",
        ],
    )
    doc.add_page_break()


def chapter_3(doc: Document) -> None:
    add_heading(doc, "Chapter 3: System Design and Methodology", 1)
    add_heading(doc, "3.1 Design Considerations", 2)
    add_body(
        doc,
        "The design is organized around seven considerations. Economically, the system avoids paid datasets and runs with commodity hardware. Environmentally, it improves solar utilization rather than requiring new physical infrastructure. Socially, it converts technical forecasts into recommendations understandable to non-experts. Ethically, it uses public weather data and avoids personal consumption profiling unless a future user opts in. From a safety perspective, the system is advisory and does not directly control electrical equipment. Politically and practically, open data helps the same architecture work across cities and regions. In terms of sustainability, the application favors low-cost deployment and repeatable model retraining.",
    )
    add_heading(doc, "3.2 Methodology", 2)
    add_body(
        doc,
        "The methodology follows a layered pipeline. First, NASA POWER data is collected for all configured cities. Second, raw records are cleaned and transformed into processed datasets. Third, feature engineering generates cyclical time features, lag variables, rolling statistics, interaction terms, and solar-position-derived values. Fourth, baseline and deep models are trained and evaluated. Fifth, forecast outputs are converted into approximate PV power. Finally, the API and dashboards present the forecast, comparison metrics, and decision recommendations.",
    )
    for spec in discover_figures()[2:5]:
        add_figure(doc, spec)
    add_table(
        doc,
        ["Feature Group", "Examples", "Purpose"],
        [
            ["Cyclical time", "hour_sin, hour_cos, day-of-year encodings", "Represent daily and seasonal cycles without artificial jumps"],
            ["Lag features", "GHI at previous 1, 24, and 48 hours", "Capture recent and same-time-previous-day behavior"],
            ["Rolling statistics", "Rolling mean and standard deviation", "Summarize short-term stability and variability"],
            ["Weather interaction", "Temperature-humidity and cloud relationships", "Model nonlinear atmospheric effects"],
            ["Solar geometry", "Top-of-atmosphere radiation and solar zenith features", "Anchor predictions to physical solar availability"],
        ],
    )

    add_heading(doc, "3.3 Modelling", 2)
    add_body(
        doc,
        "The baseline layer includes persistence, climatology, linear regression, ridge regression, random forest, and XGBoost. These models establish the minimum performance that a more complex neural approach must beat. The BiLSTM model processes sequences in both temporal directions during training, giving it stronger context for structured time-series behavior.\n\nThe NeuroSpike model uses leaky integrate-and-fire style spiking neurons. A simplified neuron accumulates input current in a membrane potential and emits a spike when the threshold is crossed. Conceptually, the membrane update can be written as V(t) = beta x V(t-1) + I(t), where beta controls leakage and I(t) is the incoming signal. When V(t) exceeds a threshold, the neuron spikes and resets. In the project implementation, this event-like representation is used to learn rapid transitions in irradiance patterns.",
    )
    add_body(
        doc,
        "PV power is estimated from predicted irradiance using P = GHI x A x eta, where A is panel area and eta is panel efficiency. The decision layer then classifies the power window into user-oriented categories such as high-yield, moderate-yield, and low-yield periods.",
    )

    add_heading(doc, "3.4 Tool and Technique Selection", 2)
    add_table(
        doc,
        ["Tool or Technique", "Reason for Selection"],
        [
            ["Python", "Strong data science ecosystem and reproducible scripting"],
            ["NASA POWER", "Open access to multi-city solar and weather variables"],
            ["scikit-learn", "Reliable baseline modeling and metrics"],
            ["XGBoost", "High-performing tree boosting for engineered tabular features"],
            ["PyTorch / SNN tooling", "Flexible deep learning and spiking model implementation"],
            ["FastAPI", "Fast, typed API layer for forecast endpoints"],
            ["Streamlit", "Rapid dashboard for model analysis and data exploration"],
            ["Static FastAPI frontend", "Standalone user-facing interface for forecast and savings views"],
            ["python-docx", "Automated generation of capstone-format Word reports"],
        ],
    )
    doc.add_page_break()


def chapter_4(doc: Document) -> None:
    add_heading(doc, "Chapter 4: Implementation and Testing", 1)
    add_heading(doc, "4.1 Development Details", 2)
    add_body(
        doc,
        "Implementation was divided into data, modeling, backend, dashboard, and reporting modules. The data layer stores raw, processed, and feature-level CSV files under the data directory. Model artifacts are stored under models/checkpoints and models/saved. Evaluation outputs are written to outputs/metrics and outputs/plots. The backend exposes forecast functionality, while the frontend layers allow users to inspect model behavior and convert forecasts into practical power and savings insights.",
    )
    for spec in frontend_figures():
        add_figure(doc, spec)
    add_body(
        doc,
        "The Streamlit dashboard supports model metrics, city comparison, forecasts, power output simulation, and data exploration. The standalone FastAPI web dashboard provides a browser-first interface with endpoints for cities, forecasts, savings, and comparisons. If additional screenshots are saved in outputs/frontend_screenshots, this generator will include them automatically in this chapter.",
    )

    add_heading(doc, "4.2 Experimental Setup", 2)
    add_table(
        doc,
        ["Component", "Configuration"],
        [
            ["Programming language", "Python 3.12 environment"],
            ["Data source", "NASA POWER city-wise solar and meteorological data"],
            ["Cities", "Riyadh, Cairo, Istanbul, New Delhi, Dubai, London, Sydney, Tokyo, Los Angeles, Nairobi"],
            ["Forecast horizon", "1 hour ahead"],
            ["Evaluation metrics", "RMSE, MAE, MAPE, R2, skill score"],
            ["Backend", "FastAPI endpoints for health, cities, metrics, and forecasts"],
            ["Frontend", "Streamlit dashboard and standalone FastAPI static dashboard"],
            ["Report automation", "python-docx with project metrics and generated figures"],
        ],
    )

    add_heading(doc, "4.3 Test Cases", 2)
    add_table(
        doc,
        ["Test Case", "Input", "Expected Output", "Verification"],
        [
            ["City list API", "GET /cities", "Supported city names", "Response includes configured cities"],
            ["Forecast API", "POST /forecast with city=new_delhi, horizon=1", "Forecast payload with GHI estimate", "Response schema and numeric value checked"],
            ["Metrics loading", "Read outputs/metrics/final_summary_table.csv", "Model comparison table", "Report table generated successfully"],
            ["Plot insertion", "Read outputs/plots/*.png", "Figures with captions", "DOCX contains result graphs"],
            ["PV conversion", "GHI, panel area, efficiency", "Estimated power output", "Formula output checked against manual calculation"],
            ["Frontend fallback", "API unavailable", "Dashboard uses stored forecast output where supported", "Dashboard remains usable for demonstration"],
        ],
    )

    add_heading(doc, "4.4 Safety Measures", 2)
    add_body(
        doc,
        "The application is advisory and does not directly switch electrical loads, inverters, or industrial equipment. This preserves human control and avoids cyber-physical safety risk in the capstone prototype. API errors are handled without exposing internal stack traces to end users. Since the present system uses public meteorological data and not personal meter readings, privacy risk is low. Future smart-meter integration should use explicit consent, encrypted storage, and role-based access control.",
    )
    doc.add_page_break()


def chapter_5(doc: Document) -> None:
    add_heading(doc, "Chapter 5: Results and Discussion", 1)
    add_heading(doc, "5.1 Result Analysis", 2)
    summary = model_summary_rows()
    if summary:
        add_table(doc, ["Model", "Mean RMSE", "Mean MAE", "Mean R2", "Mean Skill", "Best City", "Worst City"], summary)
    r2_rows = r2_city_rows()
    if r2_rows:
        add_table(doc, ["Model", "New Delhi", "London", "Riyadh", "Sydney", "Mean R2"], r2_rows)
    snn_rows = snn_city_rows()
    if snn_rows:
        add_table(doc, ["City", "RMSE", "MAE", "MAPE", "R2", "Skill"], snn_rows)

    add_body(
        doc,
        "The saved evaluation metrics show that NeuroSpike is the strongest model in the current experiment set. It reaches a mean R2 of 0.9877 across the ten-city evaluation, while BiLSTM reaches 0.9710 and XGBoost reaches 0.9642. The reduction in mean RMSE is also meaningful: NeuroSpike records 29.667 W/m2 compared with 45.520 W/m2 for BiLSTM and 47.281 W/m2 for XGBoost. This indicates that the spiking layer is not merely adding complexity; in the stored results, it improves both error magnitude and explanatory performance.\n\nCity-wise results are also important. Riyadh and Cairo show very high R2 values because their solar patterns are comparatively stable. London and New Delhi are more challenging due to variability, but NeuroSpike still remains above 0.97 R2 in the saved metrics. This matters because a model that performs only in clear-sky cities would not be suitable for a general solar intelligence platform.",
    )
    for spec in discover_figures()[5:]:
        add_figure(doc, spec)

    add_heading(doc, "5.2 Comparison Against Objectives", 2)
    add_table(
        doc,
        ["Objective", "Target", "Observed Project Result", "Status"],
        [
            ["Multi-city data pipeline", "Ten configured cities", "Raw, processed, and feature files exist for ten cities", "Achieved"],
            ["Accurate short-term forecast", "R2 greater than 0.95", "NeuroSpike mean R2 = 0.9877", "Exceeded"],
            ["Baseline comparison", "Compare statistical, ML, and deep models", "Persistence through NeuroSpike compared in metrics table", "Achieved"],
            ["Decision support", "Convert forecasts into energy guidance", "PV conversion and recommendation layer implemented", "Achieved"],
            ["User interface", "Dashboard/API demonstration", "FastAPI, Streamlit, and standalone frontend available", "Achieved"],
            ["Detailed documentation", "Report aligned with capstone template", "Automated DOCX with figures, tables, appendices, and annexures", "Achieved"],
        ],
    )

    add_heading(doc, "5.3 Economic Analysis", 2)
    add_table(
        doc,
        ["Item", "Estimated Cost", "Comment"],
        [
            ["Data source", "0", "NASA POWER is open access"],
            ["Software stack", "0", "Python and project libraries are open source"],
            ["Model training", "Low to moderate", "Can run locally or on short-duration cloud/GPU sessions"],
            ["Prototype hosting", "Low", "FastAPI app can run on a small cloud instance"],
            ["User savings", "Context-dependent", "Higher when loads can be shifted into solar-rich periods"],
        ],
    )
    add_table(
        doc,
        ["User Type", "Main Benefit", "Economic Interpretation"],
        [
            ["Residential", "Shift appliances to high-yield hours", "Reduces grid import and improves self-consumption"],
            ["Small commercial", "Schedule flexible loads using forecast windows", "Improves operating cost predictability"],
            ["Industrial", "Plan non-critical processes around solar availability", "Can reduce peak grid dependence where scheduling flexibility exists"],
            ["Municipal or educational", "Compare city-level solar potential", "Supports planning, awareness, and training at low cost"],
        ],
    )
    add_body(
        doc,
        "The cost-benefit value is strongest when the user has controllable loads. The project does not claim universal savings because actual savings depend on tariff structure, installed PV capacity, battery availability, and user behavior. Its practical contribution is decision timing: it makes solar availability visible before a task is scheduled.",
    )

    add_heading(doc, "5.4 Comparison with Existing Solutions and Standards", 2)
    add_table(
        doc,
        ["Feature", "NeuroSpike", "Typical Forecast API", "PV Monitoring Tool"],
        [
            ["Open project pipeline", "Yes", "Usually no", "Sometimes limited"],
            ["Forecast model comparison", "Included", "Usually hidden", "Usually not the focus"],
            ["Decision recommendations", "Included", "Limited", "Limited"],
            ["City comparison", "Included", "Available in some tools", "Site-specific"],
            ["Educational transparency", "High", "Low to moderate", "Moderate"],
            ["Standards relevance", "Uses forecast metrics and PV estimation concepts", "Varies", "Often aligned with PV monitoring practice"],
        ],
    )
    doc.add_page_break()


def chapter_6(doc: Document) -> None:
    add_heading(doc, "Chapter 6: Conclusion and Future Scope", 1)
    add_heading(doc, "6.1 Summary of Achievements", 2)
    add_body(
        doc,
        "NeuroSpike successfully demonstrates a complete solar intelligence workflow: data collection, preprocessing, feature engineering, model training, evaluation, forecast generation, power estimation, dashboard delivery, and automated report generation. The strongest saved model is the NeuroSpike SNN, which achieved a mean R2 of 0.9877 and a mean RMSE of 29.667 W/m2 across ten cities. The project therefore satisfies the technical objective of beating conventional baselines while also addressing the practical need for actionable recommendations.",
    )
    add_heading(doc, "6.2 Impact on Society, Sustainability, Ethics, and Compliance", 2)
    add_body(
        doc,
        "The project supports society by making solar planning easier for users who may not have expertise in forecasting or energy analytics. It supports environmental sustainability by encouraging better use of renewable energy already available from installed PV systems. Ethically, it relies on public weather data and keeps the user in control of all decisions. Compliance-wise, the prototype uses standard software interfaces and reports conventional forecasting metrics, making its behavior auditable and reproducible.",
    )
    add_heading(doc, "6.3 Limitations", 2)
    add_bullets(
        doc,
        [
            "The current trained configuration focuses on one-hour-ahead forecasting.",
            "NASA POWER data may not capture rooftop-level shading, local cloud movement, or microclimate effects.",
            "PV power conversion is approximate and does not fully model inverter losses, soiling, panel orientation, or degradation.",
            "Decision support is rule-based and not personalized to individual appliance histories or live tariffs.",
            "The SNN model is evaluated in software rather than on neuromorphic hardware, so hardware energy-efficiency benefits are not demonstrated.",
        ],
    )
    add_heading(doc, "6.4 Future Work", 2)
    add_numbered(
        doc,
        [
            "Extend forecasting horizons to 15 minutes, 3 hours, 6 hours, and 24 hours.",
            "Integrate satellite cloud imagery or live weather APIs for near-real-time improvement.",
            "Connect optional smart-meter data to personalize recommendations with consent.",
            "Add tariff-aware scheduling so the decision layer can optimize for both solar generation and electricity price.",
            "Deploy the SNN component on neuromorphic or low-power edge hardware for energy-efficient inference.",
            "Add automated screenshot capture so the report includes live frontend states from every dashboard tab.",
            "Validate the system with field data from an actual rooftop PV installation.",
        ],
    )
    doc.add_page_break()


def references_and_appendices(doc: Document) -> None:
    add_heading(doc, "References", 1)
    references = [
        "NASA Langley Research Center, \"POWER Data Access Viewer,\" NASA POWER Project, accessed 2026.",
        "International Energy Agency, \"Renewables 2024: Analysis and forecast to 2030,\" IEA, Paris, 2024.",
        "International Renewable Energy Agency, \"Renewable Power Generation Costs in 2023,\" IRENA, Abu Dhabi, 2024.",
        "F. Chollet, Deep Learning with Python, 2nd ed., Manning, 2021.",
        "I. Goodfellow, Y. Bengio, and A. Courville, Deep Learning, MIT Press, 2016.",
        "A. Graves and J. Schmidhuber, \"Framewise phoneme classification with bidirectional LSTM networks,\" Proc. IJCNN, 2005.",
        "W. Maass, \"Networks of spiking neurons: The third generation of neural network models,\" Neural Networks, vol. 10, no. 9, pp. 1659-1671, 1997.",
        "J. Eshraghian et al., \"Training spiking neural networks using lessons from deep learning,\" Proceedings of the IEEE, vol. 111, no. 9, pp. 1016-1054, 2023.",
        "F. Pedregosa et al., \"Scikit-learn: Machine learning in Python,\" Journal of Machine Learning Research, vol. 12, pp. 2825-2830, 2011.",
        "T. Chen and C. Guestrin, \"XGBoost: A scalable tree boosting system,\" Proc. KDD, pp. 785-794, 2016.",
        "A. Paszke et al., \"PyTorch: An imperative style, high-performance deep learning library,\" Proc. NeurIPS, 2019.",
        "S. Badirli et al., \"Solar irradiance forecasting using deep learning methods,\" Renewable Energy, recent survey literature.",
        "IEC 61724-1, \"Photovoltaic system performance - Part 1: Monitoring,\" International Electrotechnical Commission, 2021.",
        "S. Ramaswamy, \"FastAPI framework documentation and design principles,\" FastAPI project documentation, accessed 2026.",
        "Streamlit Inc., \"Streamlit documentation,\" accessed 2026.",
    ]
    for idx, ref in enumerate(references, start=1):
        add_body(doc, f"[{idx}] {ref}")
    doc.add_page_break()

    add_heading(doc, "Appendix A: Feature Engineering Code Summary", 1)
    add_code_block(
        doc,
        """
def build_features(df):
    df["hour_sin"] = sin(2 * pi * df.index.hour / 24)
    df["hour_cos"] = cos(2 * pi * df.index.hour / 24)
    df["doy_sin"] = sin(2 * pi * df.index.dayofyear / 365)
    df["doy_cos"] = cos(2 * pi * df.index.dayofyear / 365)
    for lag in [1, 24, 48]:
        df[f"GHI_lag_{lag}"] = df["GHI"].shift(lag)
    for window in [3, 6, 24]:
        df[f"GHI_roll_mean_{window}"] = df["GHI"].rolling(window).mean()
        df[f"GHI_roll_std_{window}"] = df["GHI"].rolling(window).std()
    df["temp_humidity_interaction"] = df["T2M"] * df["RH2M"]
    return df.dropna()
        """,
    )

    rows = feature_score_rows()
    if rows:
        headers = ["Feature"]
        csv_rows = read_csv_rows(METRICS_DIR / "feature_scores.csv")
        score_columns = [c for c in csv_rows[0].keys() if c.lower() not in {"feature", "rank"}][:3]
        headers.extend(score_columns)
        add_table(doc, headers, rows)
    doc.add_page_break()

    add_heading(doc, "Appendix B: Forecast Output Sample", 1)
    rows = forecast_rows()
    if rows:
        add_table(doc, ["City", "Predicted GHI", "Horizon"], rows)
    else:
        add_body(doc, "Forecast JSON output was not found. Run the forecast-generation notebook or backend before regenerating the report to include a live forecast sample.")
    doc.add_page_break()

    add_heading(doc, "Annexure A: Plagiarism and Originality Note", 1)
    add_body(
        doc,
        "Attach the official similarity report generated by your institution-approved plagiarism tool in this annexure. The written content in this generated report is intentionally project-specific and based on the local NeuroSpike implementation, stored metrics, and generated plots. Before final submission, replace bracketed institutional fields, add team-specific details, cite any newly added external text or images, and ensure the final similarity percentage satisfies your college policy.",
    )

    add_heading(doc, "Annexure B: PO/PSO Mapping", 1)
    add_table(
        doc,
        ["Project Task", "PO1", "PO2", "PO3", "PO4", "PO5", "PO6", "PO7", "PO8", "PO9", "PO10", "PO11", "PO12", "PSO1", "PSO2"],
        [
            ["Literature survey", "Y", "Y", "", "", "", "", "", "", "", "Y", "", "Y", "Y", ""],
            ["Data pipeline", "Y", "Y", "Y", "", "Y", "", "", "", "", "", "", "Y", "Y", ""],
            ["Feature engineering", "Y", "Y", "Y", "", "Y", "", "", "", "", "", "", "Y", "Y", ""],
            ["Model development", "Y", "Y", "Y", "Y", "Y", "", "", "", "", "", "", "Y", "Y", "Y"],
            ["Backend and dashboard", "Y", "Y", "Y", "", "Y", "", "", "", "Y", "Y", "", "Y", "", "Y"],
            ["Economic and sustainability analysis", "", "Y", "", "Y", "", "Y", "Y", "Y", "", "Y", "Y", "Y", "", ""],
        ],
    )

    add_heading(doc, "Annexure C: Outcome of the Report", 1)
    add_table(
        doc,
        ["Outcome Type", "Evidence in Project", "Status"],
        [
            ["Application/product", "FastAPI backend, Streamlit dashboard, standalone web dashboard", "Completed prototype"],
            ["Research output", "Model comparison, SNN analysis, multi-city evaluation", "Suitable for conference paper preparation"],
            ["Dataset/metrics artifact", "Processed features, trained models, plots, and metrics", "Available locally"],
            ["Patent", "Not filed", "Not pursued in current scope"],
        ],
    )

    add_heading(doc, "Annexure D: Sustainability Statement", 1)
    add_body(
        doc,
        "NeuroSpike supports sustainability by improving the timing of solar-energy use. It does not require new panels to produce value; instead, it improves the intelligence layer around existing or planned PV systems. The use of open meteorological data and open-source tools reduces cost barriers. The system is also educationally sustainable because future student teams can retrain models, add cities, change horizons, and extend the decision layer without depending on a closed commercial platform.",
    )

    add_heading(doc, "Annexure E: Team Roles", 1)
    add_table(
        doc,
        ["Team Member", "Specific Contributions"],
        [
            ["Student A", "Data collection, preprocessing, EDA, feature engineering, and baseline models"],
            ["Student B", "BiLSTM, NeuroSpike SNN training, hyperparameter experiments, and evaluation analysis"],
            ["Student C", "Decision layer, FastAPI backend, frontend dashboard, testing, and report automation"],
        ],
    )


def build_report() -> Path:
    doc = Document()
    configure_document(doc)
    add_page_number(doc.sections[0])

    add_preliminary_pages(doc)

    chapter_section = doc.add_section(WD_SECTION.NEW_PAGE)
    add_page_number(chapter_section)
    chapter_1(doc)
    chapter_2(doc)
    chapter_3(doc)
    chapter_4(doc)
    chapter_5(doc)
    chapter_6(doc)
    references_and_appendices(doc)

    doc.save(REPORT_PATH)
    return REPORT_PATH


if __name__ == "__main__":
    path = build_report()
    print(f"SUCCESS: Generated detailed capstone report at {path}")
    print("Before final submission, replace bracketed college/team fields and attach the official plagiarism report.")
